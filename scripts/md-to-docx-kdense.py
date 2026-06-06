#!/usr/bin/env python3
"""Convert k_dense_invite.md to docx (minimal markdown support)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def add_formatted_paragraph(doc, line: str):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", line)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        else:
            p.add_run(part)
    return p


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    for raw in text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Intense Quote"
                for run in p.runs:
                    run.font.name = "Menlo"
                    run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if cells and all(set(c) <= {"-", " "} for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue

        if in_table and not line.startswith("|"):
            if table_rows:
                t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                t.style = "Table Grid"
                for r, row in enumerate(table_rows):
                    for c, cell in enumerate(row):
                        t.rows[r].cells[c].text = cell.replace("`", "")
                table_rows = []
            in_table = False

        if line == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            continue
        else:
            add_formatted_paragraph(doc, line)

    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        t.style = "Table Grid"
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                t.rows[r].cells[c].text = cell.replace("`", "")

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else root / "private/agent/k_dense/k_dense_invite.md"
    out = md.with_suffix(".docx")
    convert(md, out)
